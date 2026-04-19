#!/usr/bin/env python3
"""Generate a Project+ weekly status report .docx from JSON input.

Enforces a fixed section order for week-over-week stakeholder consistency:
    Header → Executive Summary → RAG + Trend → Metrics → Accomplishments →
    Upcoming → Risks → Issues → Asks → Dependencies → Appendix (Draft Gaps).
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


TBD_PATTERN = re.compile(r"\[TBD[^\]]*\]")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TBD_ORANGE = RGBColor(0xB7, 0x47, 0x1E)
BODY = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)

RAG_COLORS = {
    "Green": "70AD47",
    "Yellow": "FFC000",
    "Red": "C00000",
}

TREND_ARROW = {
    ("Green", "Green"): ("→ stable", BODY),
    ("Green", "Yellow"): ("↓ declining", RGBColor(0xE9, 0x71, 0x32)),
    ("Green", "Red"): ("↓↓ declining", RGBColor(0xC0, 0x00, 0x00)),
    ("Yellow", "Green"): ("↑ improving", RGBColor(0x70, 0xAD, 0x47)),
    ("Yellow", "Yellow"): ("→ stable", BODY),
    ("Yellow", "Red"): ("↓ declining", RGBColor(0xC0, 0x00, 0x00)),
    ("Red", "Green"): ("↑↑ improving", RGBColor(0x70, 0xAD, 0x47)),
    ("Red", "Yellow"): ("↑ improving", RGBColor(0x70, 0xAD, 0x47)),
    ("Red", "Red"): ("→ stable", RGBColor(0xC0, 0x00, 0x00)),
}


def is_tbd(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str):
        return bool(TBD_PATTERN.search(v)) or not v.strip()
    return False


def _set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _heading(doc, text, level=1):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(16 if level == 1 else 13)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h


def _text_para(doc, text, *, size=11, italic=False, color=BODY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic or (isinstance(text, str) and bool(TBD_PATTERN.search(text)))
    if isinstance(text, str) and TBD_PATTERN.search(text):
        run.font.color.rgb = TBD_ORANGE
    else:
        run.font.color.rgb = color
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = str(text)
    run.font.size = Pt(11)
    if isinstance(text, str) and TBD_PATTERN.search(text):
        run.font.italic = True
        run.font.color.rgb = TBD_ORANGE
    else:
        run.font.color.rgb = BODY


def render_header(doc, header):
    title = doc.add_paragraph()
    title_run = title.add_run(header.get("project_name", "[TBD — no project name]"))
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Weekly Status Report")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = MUTED

    meta = doc.add_table(rows=1, cols=4)
    meta.autofit = False
    meta.style = "Light Grid Accent 1"
    cells = meta.rows[0].cells
    cells[0].text = "PM"
    cells[1].text = header.get("project_manager", "[TBD]")
    cells[2].text = "Sponsor"
    cells[3].text = header.get("sponsor", "[TBD]")

    meta2 = doc.add_table(rows=1, cols=4)
    meta2.autofit = False
    meta2.style = "Light Grid Accent 1"
    cells2 = meta2.rows[0].cells
    cells2[0].text = "Week ending"
    cells2[1].text = header.get("week_ending", "[TBD]")
    cells2[2].text = "Version"
    cells2[3].text = header.get("version", "v1")


def render_rag(doc, rag, previous_rag):
    _heading(doc, "Health (RAG)", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    c1, c2 = t.rows[0].cells
    c1.width = Inches(1.5)
    c2.width = Inches(5.0)

    rag_text = rag if rag else "[TBD — PM to assign]"
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(rag_text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if rag in RAG_COLORS else TBD_ORANGE
    if rag in RAG_COLORS:
        _set_cell_background(c1, RAG_COLORS[rag])

    # Trend
    trend_text = ""
    trend_color = MUTED
    if rag in RAG_COLORS and previous_rag in RAG_COLORS:
        trend_text, trend_color = TREND_ARROW.get((previous_rag, rag), ("", MUTED))
        trend_text = f"Trend vs last week ({previous_rag}): {trend_text}"
    elif rag in RAG_COLORS and not previous_rag:
        trend_text = "Trend: — no prior baseline"
    elif not rag:
        trend_text = "Trend: — RAG not set"
    p2 = c2.paragraphs[0]
    run2 = p2.add_run(trend_text)
    run2.font.size = Pt(11)
    run2.font.italic = True
    run2.font.color.rgb = trend_color


def render_summary(doc, summary):
    _heading(doc, "Executive Summary", level=2)
    _text_para(doc, summary or "[TBD — executive summary not stated]")


def render_metrics(doc, metrics):
    _heading(doc, "Metrics", level=2)
    if not metrics:
        _text_para(doc, "(none reported this week)", italic=True, color=MUTED)
        return
    t = doc.add_table(rows=1 + len(metrics), cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    hdr[2].text = "Target"
    for i, m in enumerate(metrics, 1):
        row = t.rows[i].cells
        row[0].text = str(m.get("name", "[TBD]"))
        row[1].text = str(m.get("value", "[TBD]"))
        row[2].text = str(m.get("target") or "")


def render_list_section(doc, title, items):
    _heading(doc, title, level=2)
    if not items:
        _text_para(doc, "(none reported this week)", italic=True, color=MUTED)
        return
    for item in items:
        _bullet(doc, item)


def render_risks(doc, risks):
    _heading(doc, "Risks", level=2)
    if not risks:
        _text_para(doc, "(none reported this week)", italic=True, color=MUTED)
        return
    for r in risks:
        text = f"{r.get('risk', '[TBD]')} — {r.get('status', 'Open')}. Mitigation: {r.get('mitigation', '[TBD]')}"
        _bullet(doc, text)


def render_issues(doc, issues):
    _heading(doc, "Issues", level=2)
    if not issues:
        _text_para(doc, "(none reported this week)", italic=True, color=MUTED)
        return
    for i in issues:
        text = f"{i.get('issue', '[TBD]')} — Owner: {i.get('owner', '[TBD]')}. Impact: {i.get('impact', '[TBD]')}"
        _bullet(doc, text)


def render_asks(doc, asks):
    _heading(doc, "Asks / Decisions Needed", level=2)
    if not asks:
        _text_para(doc, "(none this week)", italic=True, color=MUTED)
        return
    for a in asks:
        text = f"{a.get('ask', '[TBD]')} — {a.get('owner', '[TBD]')} by {a.get('needed_by', '[TBD]')}"
        _bullet(doc, text)


def render_dependencies(doc, deps):
    _heading(doc, "Dependencies", level=2)
    if not deps:
        _text_para(doc, "(none reported)", italic=True, color=MUTED)
        return
    for d in deps:
        _bullet(doc, d)


def collect_gaps(data):
    gaps = []
    header = data.get("header", {}) or {}
    for k in ("project_name", "project_manager", "sponsor", "week_ending"):
        if is_tbd(header.get(k)):
            gaps.append(f"header.{k}: {header.get(k) or '[TBD — empty]'}")
    if is_tbd(data.get("rag")):
        gaps.append(f"rag: {data.get('rag') or '[TBD — empty]'}")
    if is_tbd(data.get("executive_summary")):
        gaps.append(f"executive_summary: {data.get('executive_summary') or '[TBD — empty]'}")
    for a in data.get("asks") or []:
        for f in ("ask", "owner", "needed_by"):
            if is_tbd(a.get(f)):
                gaps.append(f"asks.{f}: {a.get(f) or '[TBD — empty]'}")
    return gaps


def render_appendix(doc, gaps):
    if not gaps:
        return
    _heading(doc, "Appendix — Draft Gaps", level=2)
    _text_para(doc, f"{len(gaps)} unresolved items — resolve before sending.",
               italic=True, color=MUTED)
    for g in gaps:
        _bullet(doc, g)


def render(data, output_path: Path, questions_threshold=5):
    doc = Document()
    # tighter margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    header = data.get("header", {}) or {}
    render_header(doc, header)
    render_summary(doc, data.get("executive_summary"))
    render_rag(doc, data.get("rag"), data.get("previous_rag"))
    render_metrics(doc, data.get("metrics") or [])
    render_list_section(doc, "Accomplishments", data.get("accomplishments") or [])
    render_list_section(doc, "Upcoming", data.get("upcoming") or [])
    render_risks(doc, data.get("risks") or [])
    render_issues(doc, data.get("issues") or [])
    render_asks(doc, data.get("asks") or [])
    render_dependencies(doc, data.get("dependencies") or [])
    gaps = collect_gaps(data)
    render_appendix(doc, gaps)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if len(gaps) >= questions_threshold:
        oq = output_path.parent / "open_questions.md"
        lines = ["# Open Questions — Status Report", "",
                 f"This draft has **{len(gaps)} unresolved items**.", ""]
        for g in gaps:
            lines.append(f"- {g}")
        oq.write_text("\n".join(lines))
        return gaps, oq
    return gaps, None


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--input", help="Path to JSON input (omit to read stdin)")
    p.add_argument("--output", required=True)
    p.add_argument("--questions-threshold", type=int, default=5)
    args = p.parse_args()

    try:
        if args.input:
            data = json.loads(Path(args.input).read_text())
        else:
            data = json.loads(sys.stdin.read())
    except Exception as e:
        print(f"ERROR: could not parse input JSON: {e}", file=sys.stderr)
        return 2

    out = Path(args.output)
    gaps, oq = render(data, out, args.questions_threshold)
    print(f"Status report written: {out}")
    print(f"TBD count: {len(gaps)}")
    if oq:
        print(f"Open questions: {oq}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
