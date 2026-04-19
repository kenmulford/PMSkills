#!/usr/bin/env python3
"""Generate a Project+ lessons-learned register .docx from JSON input.

Enforces honesty: does not balance negatives with fabricated positives and
preserves the user's exact phrasing in the what_happened field.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm


TBD_PATTERN = re.compile(r"\[TBD[^\]]*\]")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TBD_ORANGE = RGBColor(0xB7, 0x47, 0x1E)
BODY = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)
KEEP_GREEN = RGBColor(0x70, 0xAD, 0x47)
CHANGE_RED = RGBColor(0xC0, 0x00, 0x00)

VALID_CATEGORIES = {"Process", "Technical", "People", "Vendor", "Scope", "Schedule", "Cost"}
VALID_DISPOSITIONS = {"Keep", "Change"}


def is_tbd(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str):
        return bool(TBD_PATTERN.search(v)) or not v.strip()
    return False


def _heading(doc, text, level=1, color=NAVY):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(16 if level == 1 else 13)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h


def _text_para(doc, text, *, size=11, italic=False, color=BODY):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    is_tbd_text = isinstance(text, str) and bool(TBD_PATTERN.search(text))
    run.font.italic = italic or is_tbd_text
    run.font.color.rgb = TBD_ORANGE if is_tbd_text else color
    return p


def render_header(doc, header):
    title = doc.add_paragraph()
    t_run = title.add_run(f"Lessons Learned — {header.get('project_name', '[TBD]')}")
    t_run.bold = True
    t_run.font.size = Pt(20)
    t_run.font.color.rgb = NAVY

    phase = header.get("phase", "[TBD]")
    sub = doc.add_paragraph()
    s_run = sub.add_run(f"Phase: {phase}")
    s_run.font.size = Pt(12)
    s_run.font.italic = True
    s_run.font.color.rgb = MUTED

    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Light Grid Accent 1"
    rows = [
        ("PM", header.get("project_manager", "[TBD]"),
         "Sponsor", header.get("sponsor", "[TBD]")),
        ("Compiled", header.get("date_compiled", "[TBD]"),
         "Contributors", ", ".join(header.get("contributors") or []) or "[TBD]"),
    ]
    for i, (k1, v1, k2, v2) in enumerate(rows):
        cells = meta.rows[i].cells
        cells[0].text = k1
        cells[1].text = str(v1)
        cells[2].text = k2
        cells[3].text = str(v2)


def render_summary_counts(doc, lessons):
    _heading(doc, "Summary", level=2)
    keep = sum(1 for l in lessons if l.get("disposition") == "Keep")
    change = sum(1 for l in lessons if l.get("disposition") == "Change")
    tbd = sum(1 for l in lessons if l.get("disposition") not in VALID_DISPOSITIONS)
    total = len(lessons)

    p = doc.add_paragraph()
    for text, color in [
        (f"{total} lessons total  ", BODY),
        (f"• {keep} Keep  ", KEEP_GREEN),
        (f"• {change} Change  ", CHANGE_RED),
        (f"• {tbd} TBD", MUTED),
    ]:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color

    # Category breakdown
    cats = Counter()
    for l in lessons:
        c = l.get("category")
        cats[c if c in VALID_CATEGORIES else "[TBD]"] += 1
    if cats:
        parts = [f"{c}: {n}" for c, n in sorted(cats.items())]
        _text_para(doc, "By category — " + ", ".join(parts), italic=True, color=MUTED)


def _render_lesson_table(doc, lessons):
    if not lessons:
        _text_para(doc, "(none)", italic=True, color=MUTED)
        return
    t = doc.add_table(rows=1 + len(lessons), cols=6)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, label in enumerate(["ID", "Category", "What Happened", "Impact", "Recommendation", "Owner"]):
        hdr[i].text = label
    for i, l in enumerate(lessons, 1):
        cells = t.rows[i].cells
        cells[0].text = str(l.get("id", f"L-{i:02d}"))
        cells[1].text = str(l.get("category", "[TBD]"))
        cells[2].text = str(l.get("what_happened", "[TBD]"))
        cells[3].text = str(l.get("impact", "[TBD]"))
        cells[4].text = str(l.get("recommendation", "[TBD]"))
        cells[5].text = str(l.get("owner", "[TBD]"))
        # italicize TBD cells
        for c in cells:
            for para in c.paragraphs:
                for run in para.runs:
                    if TBD_PATTERN.search(run.text or ""):
                        run.font.italic = True
                        run.font.color.rgb = TBD_ORANGE


def render_keep(doc, lessons):
    keep = [l for l in lessons if l.get("disposition") == "Keep"]
    _heading(doc, "Keep Doing", level=2, color=KEEP_GREEN)
    _render_lesson_table(doc, keep)


def render_change(doc, lessons):
    change = [l for l in lessons if l.get("disposition") == "Change"]
    _heading(doc, "Change Next Time", level=2, color=CHANGE_RED)
    _render_lesson_table(doc, change)


def render_undecided(doc, lessons):
    und = [l for l in lessons if l.get("disposition") not in VALID_DISPOSITIONS]
    if not und:
        return
    _heading(doc, "Undecided (Keep or Change TBD)", level=2, color=MUTED)
    _render_lesson_table(doc, und)


def collect_gaps(data):
    gaps = []
    header = data.get("header", {}) or {}
    for k in ("project_name", "project_manager", "sponsor", "phase", "date_compiled"):
        if is_tbd(header.get(k)):
            gaps.append(f"header.{k}: {header.get(k) or '[TBD]'}")
    lessons = data.get("lessons") or []
    for l in lessons:
        lid = l.get("id", "?")
        for field in ("category", "disposition", "root_cause", "recommendation", "owner"):
            if is_tbd(l.get(field)):
                gaps.append(f"{lid}.{field}: [TBD]")
    return gaps


def render_appendix(doc, gaps):
    if not gaps:
        return
    _heading(doc, "Appendix — Draft Gaps", level=2)
    _text_para(doc, f"{len(gaps)} unresolved items.", italic=True, color=MUTED)
    for g in gaps:
        p = doc.add_paragraph(style="List Bullet")
        run = p.runs[0] if p.runs else p.add_run("")
        run.text = g
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = TBD_ORANGE


def render(data, output_path: Path, questions_threshold=8):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    render_header(doc, data.get("header", {}) or {})
    lessons = data.get("lessons") or []
    render_summary_counts(doc, lessons)
    render_keep(doc, lessons)
    render_change(doc, lessons)
    render_undecided(doc, lessons)
    gaps = collect_gaps(data)
    render_appendix(doc, gaps)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if len(gaps) >= questions_threshold:
        oq = output_path.parent / "open_questions.md"
        lines = ["# Open Questions — Lessons Learned", "",
                 f"This draft has **{len(gaps)} unresolved items**.", ""]
        for g in gaps:
            lines.append(f"- {g}")
        oq.write_text("\n".join(lines))
        return gaps, oq
    return gaps, None


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--input", help="Path to JSON input (omit to read stdin)")
    p.add_argument("--output", required=True)
    p.add_argument("--questions-threshold", type=int, default=8)
    args = p.parse_args()

    try:
        if args.input:
            data = json.loads(Path(args.input).read_text())
        else:
            data = json.loads(sys.stdin.read())
    except Exception as e:
        print(f"ERROR: could not parse input JSON: {e}", file=sys.stderr)
        return 2

    out = Path(args.output)
    gaps, oq = render(data, out, args.questions_threshold)
    print(f"Lessons learned written: {out}")
    print(f"TBD count: {len(gaps)}")
    if oq:
        print(f"Open questions: {oq}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
