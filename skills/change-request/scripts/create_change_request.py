#!/usr/bin/env python3
"""Generate a Project+ change request .docx from JSON input.

Binds baseline version movement to disposition so the artifact cannot claim a
baseline shift that wasn't explicitly approved.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


TBD_PATTERN = re.compile(r"\[TBD[^\]]*\]")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TBD_ORANGE = RGBColor(0xB7, 0x47, 0x1E)
BODY = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)

VALID_DISPOSITIONS = {
    "Submitted", "Under Review", "Approved", "Rejected", "Deferred", "Withdrawn"
}
VALID_CATEGORIES = {"Scope", "Schedule", "Cost", "Quality", "Resource"}
VALID_PRIORITIES = {"Low", "Medium", "High", "Critical"}


def is_tbd(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str):
        return bool(TBD_PATTERN.search(v)) or not v.strip()
    return False


def _heading(doc, text, level=1):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(16 if level == 1 else 13)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h


def _text_para(doc, text, *, size=11, italic=False, color=BODY):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    is_tbd_text = isinstance(text, str) and bool(TBD_PATTERN.search(text))
    run.font.italic = italic or is_tbd_text
    run.font.color.rgb = TBD_ORANGE if is_tbd_text else color
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = str(text)
    run.font.size = Pt(11)
    if isinstance(text, str) and TBD_PATTERN.search(text):
        run.font.italic = True
        run.font.color.rgb = TBD_ORANGE
    else:
        run.font.color.rgb = BODY


def render_header(doc, header, category, priority):
    title = doc.add_paragraph()
    t_run = title.add_run(f"Change Request: {header.get('cr_title', '[TBD]')}")
    t_run.bold = True
    t_run.font.size = Pt(20)
    t_run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    s_run = sub.add_run(f"{header.get('project_name', '[TBD]')} — {header.get('cr_id', '[TBD]')}")
    s_run.font.size = Pt(12)
    s_run.font.italic = True
    s_run.font.color.rgb = MUTED

    t = doc.add_table(rows=3, cols=4)
    t.style = "Light Grid Accent 1"
    t.autofit = False

    rows = [
        ("Requestor", header.get("requestor", "[TBD]"),
         "Project Manager", header.get("project_manager", "[TBD]")),
        ("Sponsor", header.get("sponsor", "[TBD]"),
         "Date Submitted", header.get("date_submitted", "[TBD]")),
        ("Category", category or "[TBD]",
         "Priority", priority or "[TBD]"),
    ]
    for i, (k1, v1, k2, v2) in enumerate(rows):
        cells = t.rows[i].cells
        cells[0].text = k1
        cells[1].text = str(v1)
        cells[2].text = k2
        cells[3].text = str(v2)


def render_description(doc, description, justification):
    _heading(doc, "Description", level=2)
    _text_para(doc, description or "[TBD — description not stated]")
    _heading(doc, "Justification", level=2)
    _text_para(doc, justification or "[TBD — justification not stated]")


def render_impact(doc, impact):
    _heading(doc, "Impact Analysis", level=2)
    t = doc.add_table(rows=6, cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False

    # Schedule — use integer if present, else verbatim estimate
    sched_days = impact.get("schedule_impact_days")
    sched_est = impact.get("schedule_user_estimate") or ""
    if sched_days is not None:
        sched_text = f"{sched_days} business days"
    elif sched_est.strip():
        sched_text = f"[TBD — user said \"{sched_est}\"]"
    else:
        sched_text = "(not stated)"

    # Cost — same pattern
    cost_usd = impact.get("cost_impact_usd")
    cost_est = impact.get("cost_user_estimate") or ""
    if cost_usd is not None:
        cost_text = f"${cost_usd:,.0f}"
    elif cost_est.strip():
        cost_text = f"[TBD — user said \"{cost_est}\"]"
    else:
        cost_text = "(not stated)"

    rows = [
        ("Scope", impact.get("scope") or "(not stated)"),
        ("Schedule", sched_text),
        ("Cost", cost_text),
        ("Resources", impact.get("resources") or "(not stated)"),
        ("Quality", impact.get("quality") or "(not stated)"),
        ("Risk", impact.get("risk") or "(not stated)"),
    ]
    for i, (k, v) in enumerate(rows):
        cells = t.rows[i].cells
        cells[0].text = k
        cells[1].text = str(v)
        # highlight TBD cells
        if isinstance(v, str) and TBD_PATTERN.search(v):
            for run in cells[1].paragraphs[0].runs:
                run.font.italic = True
                run.font.color.rgb = TBD_ORANGE


def render_options(doc, options):
    _heading(doc, "Options Considered", level=2)
    if not options:
        _text_para(doc, "(none stated)", italic=True, color=MUTED)
        return
    t = doc.add_table(rows=1 + len(options), cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Option"
    hdr[1].text = "Pros"
    hdr[2].text = "Cons"
    for i, o in enumerate(options, 1):
        cells = t.rows[i].cells
        cells[0].text = str(o.get("option", "[TBD]"))
        cells[1].text = str(o.get("pros", ""))
        cells[2].text = str(o.get("cons", ""))


def render_recommendation(doc, recommendation):
    _heading(doc, "Recommendation", level=2)
    _text_para(doc, recommendation or "[TBD — no recommendation stated]")


def render_ccb(doc, members, approvals):
    _heading(doc, "Change Control Board", level=2)
    if not members:
        _text_para(doc, "[TBD — CCB roster not stated]", italic=True)
    else:
        for m in members:
            _bullet(doc, m)

    _heading(doc, "Approvals", level=2)
    if not approvals:
        _text_para(doc, "(no decisions recorded)", italic=True, color=MUTED)
        return
    t = doc.add_table(rows=1 + len(approvals), cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Name"
    hdr[1].text = "Role"
    hdr[2].text = "Decision"
    hdr[3].text = "Date"
    for i, a in enumerate(approvals, 1):
        cells = t.rows[i].cells
        cells[0].text = str(a.get("name", "[TBD]"))
        cells[1].text = str(a.get("role", ""))
        cells[2].text = str(a.get("decision", "Pending"))
        cells[3].text = str(a.get("date", ""))


def increment_version(v: str) -> str:
    """v1.2 -> v1.3 ; v2 -> v2.1 ; fallback appends .1"""
    if not v:
        return "[TBD]"
    m = re.match(r"^v?(\d+)(?:\.(\d+))?$", v.strip(), re.IGNORECASE)
    if not m:
        return f"{v} (+1)"
    major = int(m.group(1))
    minor = int(m.group(2)) if m.group(2) is not None else 0
    return f"v{major}.{minor + 1}"


def render_disposition(doc, disposition, current_baseline):
    _heading(doc, "Disposition & Baseline", level=2)
    disp = disposition if disposition in VALID_DISPOSITIONS else "Submitted"

    if disp == "Approved":
        new_version = increment_version(current_baseline or "")
        text = f"Disposition: {disp}. Baseline moves from {current_baseline or '[TBD]'} → {new_version}."
    else:
        text = f"Disposition: {disp}. No baseline change — baseline remains {current_baseline or '[TBD]'}."

    _text_para(doc, text)


def collect_gaps(data):
    gaps = []
    header = data.get("header", {}) or {}
    for k in ("cr_id", "cr_title", "project_name", "requestor", "date_submitted", "current_baseline_version"):
        if is_tbd(header.get(k)):
            gaps.append(f"header.{k}: {header.get(k) or '[TBD — empty]'}")
    for k in ("category", "priority"):
        v = data.get(k)
        if is_tbd(v) or (isinstance(v, str) and v not in (VALID_CATEGORIES if k == "category" else VALID_PRIORITIES) and TBD_PATTERN.search(v or "")):
            gaps.append(f"{k}: {v or '[TBD]'}")
    if is_tbd(data.get("description")):
        gaps.append("description: [TBD]")
    if is_tbd(data.get("justification")):
        gaps.append("justification: [TBD]")
    impact = data.get("impact") or {}
    if impact.get("schedule_impact_days") is None and impact.get("schedule_user_estimate"):
        gaps.append(f"impact.schedule: fuzzy — \"{impact.get('schedule_user_estimate')}\"")
    if impact.get("cost_impact_usd") is None and impact.get("cost_user_estimate"):
        gaps.append(f"impact.cost: fuzzy — \"{impact.get('cost_user_estimate')}\"")
    if not (data.get("ccb_members") or []):
        gaps.append("ccb_members: [TBD — roster not stated]")
    if is_tbd(data.get("recommendation")):
        gaps.append("recommendation: [TBD]")
    return gaps


def render_integrity(doc, data):
    """Surface non-standard enum values and baseline/disposition mismatches."""
    issues = []
    disp = data.get("disposition")
    if disp and disp not in VALID_DISPOSITIONS:
        issues.append(f"disposition '{disp}' is not a standard value")
    cat = data.get("category")
    if cat and cat not in VALID_CATEGORIES and not (isinstance(cat, str) and TBD_PATTERN.search(cat)):
        issues.append(f"category '{cat}' is not a standard Project+ category")
    pri = data.get("priority")
    if pri and pri not in VALID_PRIORITIES and not (isinstance(pri, str) and TBD_PATTERN.search(pri)):
        issues.append(f"priority '{pri}' is not standard")
    # Approvals with decision=Approve but disposition != Approved
    approvals = data.get("approvals") or []
    if disp == "Approved" and not any(a.get("decision") == "Approve" for a in approvals):
        issues.append("disposition Approved but no 'Approve' decision in approvals list")
    if not issues:
        return
    _heading(doc, "Integrity Notes", level=2)
    for i in issues:
        _bullet(doc, i)


def render_appendix(doc, gaps):
    if not gaps:
        return
    _heading(doc, "Appendix — Draft Gaps", level=2)
    _text_para(doc, f"{len(gaps)} unresolved items — resolve before submitting to CCB.",
               italic=True, color=MUTED)
    for g in gaps:
        _bullet(doc, g)


def render(data, output_path: Path, questions_threshold=5):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    header = data.get("header", {}) or {}
    render_header(doc, header, data.get("category"), data.get("priority"))
    render_description(doc, data.get("description"), data.get("justification"))
    render_impact(doc, data.get("impact") or {})
    render_options(doc, data.get("options_considered") or [])
    render_recommendation(doc, data.get("recommendation"))
    render_ccb(doc, data.get("ccb_members") or [], data.get("approvals") or [])
    render_disposition(doc, data.get("disposition"), header.get("current_baseline_version"))
    render_integrity(doc, data)
    gaps = collect_gaps(data)
    render_appendix(doc, gaps)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if len(gaps) >= questions_threshold:
        oq = output_path.parent / "open_questions.md"
        lines = ["# Open Questions — Change Request", "",
                 f"This draft has **{len(gaps)} unresolved items**.", ""]
        for g in gaps:
            lines.append(f"- {g}")
        oq.write_text("\n".join(lines))
        return gaps, oq
    return gaps, None


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--input", help="Path to JSON input (omit to read stdin)")
    p.add_argument("--output", required=True)
    p.add_argument("--questions-threshold", type=int, default=5)
    args = p.parse_args()

    try:
        if args.input:
            data = json.loads(Path(args.input).read_text())
        else:
            data = json.loads(sys.stdin.read())
    except Exception as e:
        print(f"ERROR: could not parse input JSON: {e}", file=sys.stderr)
        return 2

    out = Path(args.output)
    gaps, oq = render(data, out, args.questions_threshold)
    print(f"Change request written: {out}")
    print(f"TBD count: {len(gaps)}")
    if oq:
        print(f"Open questions: {oq}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
