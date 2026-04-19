#!/usr/bin/env python3
"""
Generate a Stakeholder Engagement Plan .docx from a JSON input.

Enforces Project+ structure: stakeholder register, power/interest grid,
current-vs-desired engagement matrix, per-stakeholder communication plan.
Highlights TBDs and surfaces a Draft Gaps summary block at the top.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


TBD_PATTERN = re.compile(r"\[TBD[^\]]*\]")
TBD_COLOR = RGBColor(0xB7, 0x47, 0x1E)     # burnt orange
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F) # deep navy


ENGAGEMENT_LEVELS = ["Unaware", "Resistant", "Neutral", "Supportive", "Leading"]
QUADRANTS = ["Manage Closely", "Keep Satisfied", "Keep Informed", "Monitor"]


def is_tbd(value) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return bool(TBD_PATTERN.search(value)) or not value.strip()
    if isinstance(value, (list, dict)):
        return len(value) == 0
    return False


def collect_tbds(data: dict) -> list[str]:
    gaps: list[str] = []
    header = data.get("header", {}) or {}
    for k, v in header.items():
        if is_tbd(v):
            gaps.append(f"header.{k}: {v or '[TBD — empty]'}")
    stakeholders = data.get("stakeholders") or []
    if not stakeholders:
        gaps.append("stakeholders: (empty — no stakeholders provided)")
    for i, s in enumerate(stakeholders):
        name = s.get("name") or f"stakeholder[{i}]"
        label = name if not is_tbd(name) else f"stakeholder[{i}]"
        for field in ("name", "role_or_title", "organization", "interest", "power",
                      "interest_level", "quadrant", "current_engagement",
                      "desired_engagement", "concerns_or_risks"):
            val = s.get(field)
            if is_tbd(val):
                gaps.append(f"{label}.{field}: {val or '[TBD — empty]'}")
        comms = s.get("communication", {}) or {}
        for field in ("frequency", "channel", "message_type", "owner"):
            val = comms.get(field)
            if is_tbd(val):
                gaps.append(f"{label}.communication.{field}: {val or '[TBD — empty]'}")
    return gaps


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = HEADING_COLOR
    return p


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(str(text) if text is not None else "")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if isinstance(text, str) and TBD_PATTERN.search(text):
        run.font.color.rgb = TBD_COLOR
        run.italic = True
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = str(k)
        cell = table.cell(i, 1)
        cell.text = str(v) if v is not None else ""
        if TBD_PATTERN.search(cell.text):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = TBD_COLOR
                    run.italic = True
    return table


def add_draft_gaps_callout(doc, gaps):
    if not gaps:
        return
    add_heading(doc, f"Draft Gaps ({len(gaps)})", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "This draft contains unresolved items that need PM / sponsor input "
        "before the plan is used with actual stakeholders. They are listed below "
        "and highlighted in orange throughout the document."
    )
    run.italic = True
    run.font.size = Pt(10)
    for gap in gaps:
        bp = doc.add_paragraph(style="List Bullet")
        bp_run = bp.add_run(gap)
        bp_run.font.size = Pt(10)
        bp_run.font.color.rgb = TBD_COLOR
        bp_run.italic = True
    doc.add_paragraph()


def set_cell_tbd_style(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            if TBD_PATTERN.search(run.text or ""):
                run.font.color.rgb = TBD_COLOR
                run.italic = True


def render_register_table(doc, stakeholders):
    cols = ["Name", "Role / Title", "Organization", "Power", "Interest", "Quadrant"]
    table = doc.add_table(rows=1 + len(stakeholders), cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(cols):
        c = table.cell(0, i)
        c.text = label
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
    for i, s in enumerate(stakeholders, start=1):
        vals = [
            s.get("name", "[TBD]"),
            s.get("role_or_title", "[TBD]"),
            s.get("organization", "[TBD]"),
            s.get("power", "[TBD]"),
            s.get("interest_level", "[TBD]"),
            s.get("quadrant", "[TBD]"),
        ]
        for j, v in enumerate(vals):
            cell = table.cell(i, j)
            cell.text = str(v) if v is not None else "[TBD]"
            set_cell_tbd_style(cell)


def render_power_interest_grid(doc, stakeholders):
    """A simple 2x2 text-grid showing which stakeholders belong in which quadrant."""
    quadrant_members: dict[str, list[str]] = {q: [] for q in QUADRANTS}
    quadrant_members["[TBD]"] = []
    for s in stakeholders:
        q = s.get("quadrant") or "[TBD]"
        if is_tbd(q):
            quadrant_members["[TBD]"].append(s.get("name", "[TBD]"))
        elif q in quadrant_members:
            quadrant_members[q].append(s.get("name", "[TBD]"))
        else:
            quadrant_members.setdefault(q, []).append(s.get("name", "[TBD]"))

    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"
    # Top-left header
    table.cell(0, 0).text = ""
    table.cell(0, 1).text = "Low Interest"
    table.cell(0, 2).text = "High Interest"
    table.cell(1, 0).text = "High Power"
    table.cell(2, 0).text = "Low Power"
    table.cell(1, 1).text = "Keep Satisfied:\n" + ("\n".join(quadrant_members["Keep Satisfied"]) or "—")
    table.cell(1, 2).text = "Manage Closely:\n" + ("\n".join(quadrant_members["Manage Closely"]) or "—")
    table.cell(2, 1).text = "Monitor:\n" + ("\n".join(quadrant_members["Monitor"]) or "—")
    table.cell(2, 2).text = "Keep Informed:\n" + ("\n".join(quadrant_members["Keep Informed"]) or "—")

    for cell in (table.cell(0, 1), table.cell(0, 2), table.cell(1, 0), table.cell(2, 0)):
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    if quadrant_members["[TBD]"]:
        p = doc.add_paragraph()
        run = p.add_run(
            f"Unplaced (quadrant TBD): {', '.join(quadrant_members['[TBD]'])}"
        )
        run.italic = True
        run.font.color.rgb = TBD_COLOR
        run.font.size = Pt(10)


def render_engagement_matrix(doc, stakeholders):
    cols = ["Stakeholder"] + ENGAGEMENT_LEVELS
    table = doc.add_table(rows=1 + len(stakeholders), cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(cols):
        c = table.cell(0, i)
        c.text = label
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
    for i, s in enumerate(stakeholders, start=1):
        name = s.get("name", "[TBD]")
        current = s.get("current_engagement")
        desired = s.get("desired_engagement")
        table.cell(i, 0).text = str(name)
        for j, level in enumerate(ENGAGEMENT_LEVELS, start=1):
            marks = []
            if isinstance(current, str) and current.strip() == level:
                marks.append("C")
            if isinstance(desired, str) and desired.strip() == level:
                marks.append("D")
            table.cell(i, j).text = " / ".join(marks)
        # If both current and desired are TBD, add a TBD cell note
        if is_tbd(current) and is_tbd(desired):
            # overwrite with a TBD marker in last cell
            table.cell(i, len(ENGAGEMENT_LEVELS)).text = "[TBD]"
            set_cell_tbd_style(table.cell(i, len(ENGAGEMENT_LEVELS)))

    p = doc.add_paragraph()
    run = p.add_run("Legend: C = Current engagement level, D = Desired engagement level. "
                    "A gap between C and D indicates stakeholders who need active attention.")
    run.italic = True
    run.font.size = Pt(9)


def render_comms_plan(doc, stakeholders):
    cols = ["Stakeholder", "Frequency", "Channel", "Message Type", "Owner"]
    table = doc.add_table(rows=1 + len(stakeholders), cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, label in enumerate(cols):
        c = table.cell(0, i)
        c.text = label
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
    for i, s in enumerate(stakeholders, start=1):
        comms = s.get("communication", {}) or {}
        vals = [
            s.get("name", "[TBD]"),
            comms.get("frequency", "[TBD]"),
            comms.get("channel", "[TBD]"),
            comms.get("message_type", "[TBD]"),
            comms.get("owner", "[TBD]"),
        ]
        for j, v in enumerate(vals):
            cell = table.cell(i, j)
            cell.text = str(v) if v is not None else "[TBD]"
            set_cell_tbd_style(cell)


def render_sep(data: dict, output_path: Path) -> list[str]:
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    header = data.get("header", {}) or {}
    project_name = header.get("project_name", "[TBD — no project name]")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{project_name} — Stakeholder Engagement Plan")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = HEADING_COLOR

    add_kv_table(doc, [
        ("Project Name", header.get("project_name", "[TBD]")),
        ("Project Manager", header.get("project_manager", "[TBD]")),
        ("Sponsor", header.get("sponsor", "[TBD]")),
        ("Date", header.get("date", "[TBD]")),
        ("Version", header.get("version", "v1")),
    ])
    doc.add_paragraph()

    gaps = collect_tbds(data)
    add_draft_gaps_callout(doc, gaps)

    stakeholders = data.get("stakeholders") or []

    add_heading(doc, "1. Stakeholder Register", level=1)
    if not stakeholders:
        add_para(doc, "[TBD — no stakeholders identified]")
    else:
        render_register_table(doc, stakeholders)
    doc.add_paragraph()

    add_heading(doc, "2. Power / Interest Grid", level=1)
    if stakeholders:
        render_power_interest_grid(doc, stakeholders)
    else:
        add_para(doc, "[TBD — no stakeholders to place on the grid]")
    doc.add_paragraph()

    add_heading(doc, "3. Engagement Level Matrix", level=1)
    add_para(doc,
             "Current vs. desired engagement level, per Project+ ladder "
             "(Unaware → Resistant → Neutral → Supportive → Leading).")
    if stakeholders:
        render_engagement_matrix(doc, stakeholders)
    else:
        add_para(doc, "[TBD — no stakeholders]")
    doc.add_paragraph()

    add_heading(doc, "4. Stakeholder Detail", level=1)
    if not stakeholders:
        add_para(doc, "[TBD — no stakeholder detail to show]")
    for s in stakeholders:
        add_heading(doc, s.get("name", "[TBD]"), level=2)
        add_kv_table(doc, [
            ("Role / Title", s.get("role_or_title", "[TBD]")),
            ("Organization", s.get("organization", "[TBD]")),
            ("Interest", s.get("interest", "[TBD]")),
            ("Power", s.get("power", "[TBD]")),
            ("Interest level", s.get("interest_level", "[TBD]")),
            ("Quadrant", s.get("quadrant", "[TBD]")),
            ("Current engagement", s.get("current_engagement", "[TBD]")),
            ("Desired engagement", s.get("desired_engagement", "[TBD]")),
            ("Concerns / risks", s.get("concerns_or_risks", "[TBD]")),
        ])
        doc.add_paragraph()

    add_heading(doc, "5. Communication Plan", level=1)
    if stakeholders:
        render_comms_plan(doc, stakeholders)
    else:
        add_para(doc, "[TBD — no stakeholders]")
    doc.add_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return gaps


def write_open_questions(gaps, output_path: Path):
    lines = [
        "# Open Questions for Sponsor / PM",
        "",
        f"This Stakeholder Engagement Plan draft has **{len(gaps)} unresolved items**. "
        "Resolve these before using the plan.",
        "",
    ]
    by_section: dict[str, list[str]] = {}
    for g in gaps:
        section = g.split(".")[0] or "general"
        by_section.setdefault(section, []).append(g)
    for section, items in sorted(by_section.items()):
        lines.append(f"## {section}")
        lines.append("")
        for item in items:
            lines.append(f"- {item}")
        lines.append("")
    output_path.write_text("\n".join(lines))


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a Stakeholder Engagement Plan .docx")
    parser.add_argument("--input", help="Path to JSON input (omit to read stdin)")
    parser.add_argument("--output", required=True)
    parser.add_argument("--questions-threshold", type=int, default=5)
    args = parser.parse_args()

    output_path = Path(args.output)

    try:
        if args.input:
            data = json.loads(Path(args.input).read_text())
        else:
            data = json.loads(sys.stdin.read())
    except Exception as e:
        print(f"ERROR: could not parse input JSON: {e}", file=sys.stderr)
        return 2

    gaps = render_sep(data, output_path)
    print(f"SEP written: {output_path}")
    print(f"TBD count: {len(gaps)}")

    if len(gaps) >= args.questions_threshold:
        oq = output_path.parent / "open_questions.md"
        write_open_questions(gaps, oq)
        print(f"Open questions written: {oq}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
