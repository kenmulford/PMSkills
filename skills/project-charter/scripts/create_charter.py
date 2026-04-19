#!/usr/bin/env python3
"""
Generate a project charter .docx from a JSON input.

Enforces CompTIA Project+ structure, highlights TBDs, emits a companion
open_questions.md when there are many gaps. The skill calls this script
instead of writing python-docx code from scratch.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


TBD_PATTERN = re.compile(r"\[TBD[^\]]*\]")
TBD_COLOR = RGBColor(0xB7, 0x47, 0x1E)  # burnt orange
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)  # deep navy


def is_tbd(value) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return bool(TBD_PATTERN.search(value)) or not value.strip()
    if isinstance(value, (list, dict)):
        return len(value) == 0
    return False


def collect_tbds(data: dict, prefix: str = "") -> list[str]:
    """Walk the schema and return a list of human-readable TBD locations."""
    gaps: list[str] = []

    def walk(node, path):
        if isinstance(node, dict):
            for k, v in node.items():
                walk(v, f"{path}.{k}" if path else k)
        elif isinstance(node, list):
            if not node:
                gaps.append(f"{path} (empty)")
            for i, item in enumerate(node):
                walk(item, f"{path}[{i}]")
        else:
            if isinstance(node, str) and TBD_PATTERN.search(node):
                gaps.append(f"{path}: {node}")

    walk(data, prefix)
    return gaps


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = HEADING_COLOR


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text or "")
    run.font.size = Pt(11)
    run.bold = bold
    if isinstance(text, str) and TBD_PATTERN.search(text):
        run.font.color.rgb = TBD_COLOR
        run.italic = True


def add_bullets(doc: Document, items: list[str]) -> None:
    if not items:
        items = ["[TBD — none listed]"]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(str(item))
        run.font.size = Pt(11)
        if TBD_PATTERN.search(str(item)):
            run.font.color.rgb = TBD_COLOR
            run.italic = True


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        cell = table.cell(i, 1)
        cell.text = str(v) if v is not None else ""
        for para in cell.paragraphs:
            for run in para.runs:
                if TBD_PATTERN.search(cell.text):
                    run.font.color.rgb = TBD_COLOR
                    run.italic = True


def add_draft_gaps_callout(doc: Document, gaps: list[str]) -> None:
    """Shaded summary block at the top listing all TBDs."""
    if not gaps:
        return
    add_heading(doc, f"Draft Gaps ({len(gaps)})", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "This draft contains unresolved items that need sponsor input before approval. "
        "They are listed below and highlighted in orange throughout the document."
    )
    run.italic = True
    run.font.size = Pt(10)
    for gap in gaps:
        bp = doc.add_paragraph(style="List Bullet")
        bp_run = bp.add_run(gap)
        bp_run.font.size = Pt(10)
        bp_run.font.color.rgb = TBD_COLOR
        bp_run.italic = True
    doc.add_paragraph()  # spacer


def render_charter(data: dict, output_path: Path) -> list[str]:
    doc = Document()
    # Margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    header = data.get("header", {})
    project_name = header.get("project_name", "[TBD — no project name]")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{project_name} — Project Charter")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = HEADING_COLOR

    # Header table
    add_kv_table(
        doc,
        [
            ("Project Name", header.get("project_name", "[TBD]")),
            ("Project Manager", header.get("project_manager", "[TBD]")),
            ("Sponsor", header.get("sponsor", "[TBD]")),
            ("Date", header.get("date", "[TBD]")),
            ("Version", header.get("version", "v1")),
        ],
    )
    doc.add_paragraph()

    # Draft gaps callout
    gaps = collect_tbds(data)
    add_draft_gaps_callout(doc, gaps)

    # 1. Vision
    add_heading(doc, "1. Project Vision", level=1)
    add_para(doc, data.get("vision", "[TBD — no vision stated]"))

    # 2. Objectives
    add_heading(doc, "2. Objectives and Success Measures", level=1)
    objectives = data.get("objectives") or []
    if not objectives:
        add_para(doc, "[TBD — no objectives stated]")
    for i, obj in enumerate(objectives, 1):
        add_para(doc, f"Objective {i}: {obj.get('statement', '[TBD]')}", bold=True)
        add_para(doc, f"  Leading indicator: {obj.get('leading_indicator', '[TBD]')}")
        add_para(doc, f"  Lagging indicator: {obj.get('lagging_indicator', '[TBD]')}")

    # 3. Scope Summary
    add_heading(doc, "3. Scope Summary", level=1)
    scope = data.get("scope", {})
    add_para(doc, scope.get("summary", "[TBD — no scope summary]"))
    add_para(doc, "In scope:", bold=True)
    add_bullets(doc, scope.get("in_scope") or [])
    add_para(doc, "Out of scope:", bold=True)
    out_of_scope = scope.get("out_of_scope") or []
    if not out_of_scope:
        out_of_scope = ["[TBD — ask sponsor what common scope adjacencies should be excluded]"]
    add_bullets(doc, out_of_scope)

    # 4. Preliminary Scope Statement
    add_heading(doc, "4. Preliminary Scope Statement", level=1)
    pscope = data.get("preliminary_scope", {})

    add_heading(doc, "4.1 Deliverables", level=2)
    add_bullets(doc, pscope.get("deliverables") or [])

    add_heading(doc, "4.2 Acceptance Criteria", level=2)
    add_bullets(doc, pscope.get("acceptance_criteria") or [])

    add_heading(doc, "4.3 Assumptions", level=2)
    add_bullets(doc, pscope.get("assumptions") or [])

    add_heading(doc, "4.4 Constraints", level=2)
    constraints = pscope.get("constraints", {}) or {}
    add_kv_table(
        doc,
        [
            ("Budget", constraints.get("budget", "[TBD]")),
            ("Timeline", constraints.get("timeline", "[TBD]")),
            ("Resources", constraints.get("resources", "[TBD]")),
            ("Technical / Regulatory", constraints.get("technical_or_regulatory", "[TBD]")),
        ],
    )

    # 5. Stakeholders
    add_heading(doc, "5. Stakeholders and Organization", level=1)
    stakeholders = data.get("stakeholders") or []
    if not stakeholders:
        add_para(doc, "[TBD — no stakeholders identified]")
    else:
        table = doc.add_table(rows=1 + len(stakeholders), cols=2)
        table.style = "Light Grid Accent 1"
        table.cell(0, 0).text = "Name / Group"
        table.cell(0, 1).text = "Role"
        for i, s in enumerate(stakeholders, 1):
            table.cell(i, 0).text = s.get("name_or_group", "[TBD]")
            table.cell(i, 1).text = s.get("role", "[TBD]")

    # 6. Implementation Overview
    add_heading(doc, "6. Implementation Overview", level=1)
    impl = data.get("implementation", {}) or {}

    add_heading(doc, "6.1 Milestones", level=2)
    milestones = impl.get("milestones") or []
    if not milestones:
        add_para(doc, "[TBD — no milestones listed]")
    else:
        table = doc.add_table(rows=1 + len(milestones), cols=2)
        table.style = "Light Grid Accent 1"
        table.cell(0, 0).text = "Milestone"
        table.cell(0, 1).text = "Target Date"
        for i, m in enumerate(milestones, 1):
            table.cell(i, 0).text = m.get("name", "[TBD]")
            table.cell(i, 1).text = m.get("target_date", "[TBD]")

    add_heading(doc, "6.2 Top Risks", level=2)
    risks = impl.get("risks") or []
    if not risks:
        add_para(doc, "[TBD — no risks listed]")
    else:
        table = doc.add_table(rows=1 + len(risks), cols=4)
        table.style = "Light Grid Accent 1"
        table.cell(0, 0).text = "Risk"
        table.cell(0, 1).text = "Impact"
        table.cell(0, 2).text = "Likelihood"
        table.cell(0, 3).text = "Initial Response"
        for i, r in enumerate(risks, 1):
            table.cell(i, 0).text = r.get("risk", "[TBD]")
            table.cell(i, 1).text = r.get("impact", "[TBD]")
            table.cell(i, 2).text = r.get("likelihood", "[TBD]")
            table.cell(i, 3).text = r.get("response", "[TBD]")

    add_heading(doc, "6.3 Dependencies", level=2)
    add_bullets(doc, impl.get("dependencies") or [])

    # 7. Approval
    add_heading(doc, "7. Approval", level=1)
    add_kv_table(
        doc,
        [
            (f"Project Sponsor — {header.get('sponsor', '[TBD]')}", "Signature: __________  Date: __________"),
            (f"Project Manager — {header.get('project_manager', '[TBD]')}", "Signature: __________  Date: __________"),
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return gaps


def write_open_questions(gaps: list[str], output_path: Path) -> None:
    lines = [
        "# Open Questions for Sponsor",
        "",
        f"This charter draft has **{len(gaps)} unresolved items**. Resolve these before seeking approval.",
        "",
    ]
    # Group by top-level section
    by_section: dict[str, list[str]] = {}
    for g in gaps:
        section = g.split(".")[0].split("[")[0].strip() or "general"
        by_section.setdefault(section, []).append(g)
    for section, items in sorted(by_section.items()):
        lines.append(f"## {section}")
        lines.append("")
        for item in items:
            lines.append(f"- {item}")
        lines.append("")
    output_path.write_text("\n".join(lines))


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a project charter .docx")
    parser.add_argument("--input", help="Path to JSON input file (omit to read stdin)")
    parser.add_argument("--output", required=True, help="Path to output .docx file")
    parser.add_argument(
        "--questions-threshold",
        type=int,
        default=5,
        help="Emit open_questions.md when TBD count is >= this (default 5)",
    )
    args = parser.parse_args()

    output_path = Path(args.output)

    try:
        if args.input:
            data = json.loads(Path(args.input).read_text())
        else:
            data = json.loads(sys.stdin.read())
    except Exception as e:
        print(f"ERROR: could not parse input JSON: {e}", file=sys.stderr)
        return 2

    gaps = render_charter(data, output_path)
    print(f"Charter written: {output_path}")
    print(f"TBD count: {len(gaps)}")

    if len(gaps) >= args.questions_threshold:
        oq_path = output_path.parent / "open_questions.md"
        write_open_questions(gaps, oq_path)
        print(f"Open questions written: {oq_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
