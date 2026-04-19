#!/usr/bin/env python3
"""Generate meeting minutes .docx from JSON input.

Refuses to fabricate attendees, decisions, or action items.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm


TBD_PATTERN = re.compile(r"\[TBD[^\]]*\]")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TBD_ORANGE = RGBColor(0xB7, 0x47, 0x1E)
BODY = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)


def is_tbd(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str):
        return bool(TBD_PATTERN.search(v)) or not v.strip()
    return False


def _heading(doc, text, level=1):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(16 if level == 1 else 13)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)


def _text_para(doc, text, *, size=11, italic=False, color=BODY):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    is_tbd_text = isinstance(text, str) and bool(TBD_PATTERN.search(text))
    run.font.italic = italic or is_tbd_text
    run.font.color.rgb = TBD_ORANGE if is_tbd_text else color


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = str(text)
    run.font.size = Pt(11)
    if isinstance(text, str) and TBD_PATTERN.search(text):
        run.font.italic = True
        run.font.color.rgb = TBD_ORANGE
    else:
        run.font.color.rgb = BODY


def render_header(doc, header):
    title = doc.add_paragraph()
    t_run = title.add_run(header.get("meeting_title") or "Meeting Minutes")
    t_run.bold = True
    t_run.font.size = Pt(20)
    t_run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    s_run = sub.add_run(header.get("project_name", ""))
    s_run.font.size = Pt(12)
    s_run.font.italic = True
    s_run.font.color.rgb = MUTED

    t = doc.add_table(rows=2, cols=4)
    t.style = "Light Grid Accent 1"
    rows = [
        ("Date", header.get("meeting_date", "[TBD]"),
         "Time", header.get("meeting_time") or "[TBD]"),
        ("Chair", header.get("chair", "[TBD]"),
         "Scribe", header.get("scribe", "[TBD]")),
    ]
    for i, (k1, v1, k2, v2) in enumerate(rows):
        cells = t.rows[i].cells
        cells[0].text = k1
        cells[1].text = str(v1)
        cells[2].text = k2
        cells[3].text = str(v2)

    loc = header.get("location")
    if loc:
        _text_para(doc, f"Location: {loc}", italic=True, color=MUTED)


def render_attendees(doc, attendees):
    _heading(doc, "Attendees", level=2)
    att = attendees or {}
    present = att.get("present") or []
    absent = att.get("absent") or []

    p = doc.add_paragraph()
    r = p.add_run("Present: ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(", ".join(present) if present else "[TBD — no attendees named]")
    r2.font.size = Pt(11)
    if not present:
        r2.font.italic = True
        r2.font.color.rgb = TBD_ORANGE

    p2 = doc.add_paragraph()
    r3 = p2.add_run("Absent: ")
    r3.bold = True
    r3.font.size = Pt(11)
    r4 = p2.add_run(", ".join(absent) if absent else "(none noted)")
    r4.font.size = Pt(11)
    if not absent:
        r4.font.italic = True
        r4.font.color.rgb = MUTED


def render_agenda(doc, items):
    _heading(doc, "Agenda & Discussion", level=2)
    if not items:
        _text_para(doc, "(no agenda items recorded)", italic=True, color=MUTED)
        return
    for i, item in enumerate(items, 1):
        topic = item.get("topic", "[TBD]")
        presenter = item.get("presenter") or ""
        label = f"{i}. {topic}" + (f" — {presenter}" if presenter else "")
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        disc = item.get("discussion") or ""
        if disc:
            _text_para(doc, disc, size=11)


def render_decisions(doc, decisions):
    _heading(doc, "Decisions", level=2)
    if not decisions:
        _text_para(doc, "(no decisions recorded)", italic=True, color=MUTED)
        return
    for d in decisions:
        text = d.get("decision", "[TBD]")
        by = d.get("decided_by") or "[TBD]"
        rationale = d.get("rationale") or ""
        line = f"{text} — Decided by: {by}"
        if rationale:
            line += f". Rationale: {rationale}"
        _bullet(doc, line)


def render_action_items(doc, actions):
    _heading(doc, "Action Items", level=2)
    if not actions:
        _text_para(doc, "(no action items)", italic=True, color=MUTED)
        return
    t = doc.add_table(rows=1 + len(actions), cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, label in enumerate(["Action", "Owner", "Due", "Status"]):
        hdr[i].text = label
    for i, a in enumerate(actions, 1):
        cells = t.rows[i].cells
        cells[0].text = str(a.get("action", "[TBD]"))
        cells[1].text = str(a.get("owner", "[TBD]"))
        cells[2].text = str(a.get("due_date", "[TBD]"))
        cells[3].text = str(a.get("status", "Open"))
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                if TBD_PATTERN.search(run.text or ""):
                    run.font.italic = True
                    run.font.color.rgb = TBD_ORANGE


def render_parking_lot(doc, items):
    _heading(doc, "Parking Lot", level=2)
    if not items:
        _text_para(doc, "(empty)", italic=True, color=MUTED)
        return
    for item in items:
        _bullet(doc, item)


def render_next_meeting(doc, next_meeting):
    nm = next_meeting or {}
    date_str = nm.get("date") or "[TBD]"
    topic = nm.get("topic") or ""
    _heading(doc, "Next Meeting", level=2)
    line = f"Date: {date_str}"
    if topic:
        line += f"  |  Topic: {topic}"
    _text_para(doc, line)


def collect_gaps(data):
    gaps = []
    header = data.get("header", {}) or {}
    for k in ("project_name", "meeting_title", "meeting_date", "chair"):
        if is_tbd(header.get(k)):
            gaps.append(f"header.{k}: [TBD]")
    att = data.get("attendees") or {}
    if not (att.get("present") or []):
        gaps.append("attendees.present: [TBD — no attendees]")
    for a in data.get("action_items") or []:
        if is_tbd(a.get("owner")):
            gaps.append(f"action '{a.get('action', '?')[:40]}' owner: [TBD]")
        if is_tbd(a.get("due_date")):
            gaps.append(f"action '{a.get('action', '?')[:40]}' due: [TBD]")
    for d in data.get("decisions") or []:
        if is_tbd(d.get("decided_by")):
            gaps.append(f"decision '{d.get('decision', '?')[:40]}' decided_by: [TBD]")
    nm = data.get("next_meeting") or {}
    if is_tbd(nm.get("date")):
        gaps.append("next_meeting.date: [TBD]")
    return gaps


def render_appendix(doc, gaps):
    if not gaps:
        return
    _heading(doc, "Appendix — Draft Gaps", level=2)
    _text_para(doc, f"{len(gaps)} unresolved items.", italic=True, color=MUTED)
    for g in gaps:
        _bullet(doc, g)


def render(data, output_path: Path, questions_threshold=6):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    render_header(doc, data.get("header", {}) or {})
    render_attendees(doc, data.get("attendees") or {})
    render_agenda(doc, data.get("agenda_items") or [])
    render_decisions(doc, data.get("decisions") or [])
    render_action_items(doc, data.get("action_items") or [])
    render_parking_lot(doc, data.get("parking_lot") or [])
    render_next_meeting(doc, data.get("next_meeting") or {})
    gaps = collect_gaps(data)
    render_appendix(doc, gaps)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if len(gaps) >= questions_threshold:
        oq = output_path.parent / "open_questions.md"
        lines = ["# Open Questions — Meeting Minutes", "",
                 f"This draft has **{len(gaps)} unresolved items**.", ""]
        for g in gaps:
            lines.append(f"- {g}")
        oq.write_text("\n".join(lines))
        return gaps, oq
    return gaps, None


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--input", help="Path to JSON input (omit to read stdin)")
    p.add_argument("--output", required=True)
    p.add_argument("--questions-threshold", type=int, default=6)
    args = p.parse_args()

    try:
        if args.input:
            data = json.loads(Path(args.input).read_text())
        else:
            data = json.loads(sys.stdin.read())
    except Exception as e:
        print(f"ERROR: could not parse input JSON: {e}", file=sys.stderr)
        return 2

    out = Path(args.output)
    gaps, oq = render(data, out, args.questions_threshold)
    print(f"Minutes written: {out}")
    print(f"TBD count: {len(gaps)}")
    if oq:
        print(f"Open questions: {oq}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
