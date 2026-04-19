#!/usr/bin/env python3
"""Generate a Project+ project closeout report .docx from JSON input.

Refuses to compute variance from missing data and refuses to stamp
acceptance or sign-off without explicit user input.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm


TBD_PATTERN = re.compile(r"\[TBD[^\]]*\]")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TBD_ORANGE = RGBColor(0xB7, 0x47, 0x1E)
BODY = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x70, 0xAD, 0x47)
RED = RGBColor(0xC0, 0x00, 0x00)

VALID_ACCEPTANCE = {"Accepted", "Conditionally Accepted", "Rejected", "Pending"}
VALID_SIGNOFF = {"Signed", "Pending"}


def is_tbd(v) -> bool:
    if v is None:
        return True
    if isinstance(v, str):
        return bool(TBD_PATTERN.search(v)) or not v.strip()
    return False


def _heading(doc, text, level=1, color=NAVY):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(16 if level == 1 else 13)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h


def _text_para(doc, text, *, size=11, italic=False, color=BODY):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    is_tbd_text = isinstance(text, str) and bool(TBD_PATTERN.search(text))
    run.font.italic = italic or is_tbd_text
    run.font.color.rgb = TBD_ORANGE if is_tbd_text else color
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = str(text)
    run.font.size = Pt(11)
    if isinstance(text, str) and TBD_PATTERN.search(text):
        run.font.italic = True
        run.font.color.rgb = TBD_ORANGE
    else:
        run.font.color.rgb = BODY


def render_header(doc, header):
    title = doc.add_paragraph()
    t_run = title.add_run(f"Project Closeout Report — {header.get('project_name', '[TBD]')}")
    t_run.bold = True
    t_run.font.size = Pt(20)
    t_run.font.color.rgb = NAVY

    meta = doc.add_table(rows=3, cols=4)
    meta.style = "Light Grid Accent 1"
    rows = [
        ("PM", header.get("project_manager", "[TBD]"),
         "Sponsor", header.get("sponsor", "[TBD]")),
        ("Customer", header.get("customer", "[TBD]"),
         "Closeout Date", header.get("closeout_date", "[TBD]")),
        ("Start Date", header.get("start_date", "[TBD]"),
         "", ""),
    ]
    for i, (k1, v1, k2, v2) in enumerate(rows):
        cells = meta.rows[i].cells
        cells[0].text = k1
        cells[1].text = str(v1)
        cells[2].text = k2
        cells[3].text = str(v2)


def render_summary(doc, summary):
    _heading(doc, "Executive Summary", level=2)
    _text_para(doc, summary or "[TBD — executive summary not stated]")


def render_scope(doc, delivered, deferred):
    _heading(doc, "Scope Delivered", level=2)
    if not delivered:
        _text_para(doc, "[TBD — no scope delivered listed]", italic=True)
    else:
        for item in delivered:
            _bullet(doc, item)
    if deferred:
        _heading(doc, "Scope Deferred", level=2)
        for item in deferred:
            _bullet(doc, item)


def _parse_date(s):
    if not s:
        return None
    try:
        y, m, d = s.split("-")
        return date(int(y), int(m), int(d))
    except Exception:
        return None


def render_schedule(doc, schedule):
    _heading(doc, "Schedule", level=2)
    sched = schedule or {}
    t = doc.add_table(rows=3, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = ""
    hdr[1].text = "Planned"
    hdr[2].text = "Actual"
    t.rows[1].cells[0].text = "Start"
    t.rows[1].cells[1].text = sched.get("planned_start") or "[TBD]"
    t.rows[1].cells[2].text = sched.get("actual_start") or "[TBD]"
    t.rows[2].cells[0].text = "End"
    t.rows[2].cells[1].text = sched.get("planned_end") or "[TBD]"
    t.rows[2].cells[2].text = sched.get("actual_end") or "[TBD]"

    pe = _parse_date(sched.get("planned_end"))
    ae = _parse_date(sched.get("actual_end"))
    if pe and ae:
        diff = (ae - pe).days
        if diff == 0:
            text, color = "Schedule variance: 0 days (on schedule)", GREEN
        elif diff > 0:
            text, color = f"Schedule variance: +{diff} days late", RED
        else:
            text, color = f"Schedule variance: {diff} days ahead", GREEN
    else:
        text, color = "Schedule variance: [TBD — planned_end and/or actual_end not provided]", TBD_ORANGE
    _text_para(doc, text, italic=True, color=color)


def render_cost(doc, cost):
    _heading(doc, "Cost", level=2)
    c = cost or {}
    planned = c.get("planned_usd")
    actual = c.get("actual_usd")
    currency = c.get("currency", "USD")

    t = doc.add_table(rows=2, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Planned"
    hdr[1].text = "Actual"
    hdr[2].text = "Variance"

    planned_txt = f"${planned:,.0f}" if planned is not None else "[TBD]"
    actual_txt = f"${actual:,.0f}" if actual is not None else "[TBD]"
    if planned is not None and actual is not None:
        var = planned - actual
        sign = "under" if var >= 0 else "over"
        var_txt = f"${abs(var):,.0f} {sign}"
    else:
        var_txt = "[TBD]"

    row = t.rows[1].cells
    row[0].text = planned_txt
    row[1].text = actual_txt
    row[2].text = var_txt
    # italicize TBD cells
    for cell in row:
        for run in cell.paragraphs[0].runs:
            if TBD_PATTERN.search(run.text or ""):
                run.font.italic = True
                run.font.color.rgb = TBD_ORANGE


def render_deliverables(doc, deliverables):
    _heading(doc, "Deliverables & Acceptance", level=2)
    if not deliverables:
        _text_para(doc, "[TBD — no deliverables listed]", italic=True)
        return
    t = doc.add_table(rows=1 + len(deliverables), cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, label in enumerate(["Deliverable", "Status", "Accepted By", "Date"]):
        hdr[i].text = label
    for i, d in enumerate(deliverables, 1):
        cells = t.rows[i].cells
        cells[0].text = str(d.get("name", "[TBD]"))
        cells[1].text = str(d.get("acceptance_status", "[TBD]"))
        cells[2].text = str(d.get("accepted_by", "[TBD]"))
        cells[3].text = str(d.get("acceptance_date", ""))
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                if TBD_PATTERN.search(run.text or ""):
                    run.font.italic = True
                    run.font.color.rgb = TBD_ORANGE


def render_open_items(doc, items):
    _heading(doc, "Open Items", level=2)
    if not items:
        _text_para(doc, "(none reported)", italic=True, color=MUTED)
        return
    for item in items:
        _bullet(doc, item)


def render_defects(doc, defects):
    _heading(doc, "Known Defects", level=2)
    if not defects:
        _text_para(doc, "(none reported)", italic=True, color=MUTED)
        return
    for d in defects:
        text = f"{d.get('defect', '[TBD]')} — Severity: {d.get('severity', '[TBD]')}. Workaround: {d.get('workaround', '[TBD]')}"
        _bullet(doc, text)


def render_vendor_closure(doc, vendors):
    _heading(doc, "Vendor / Contract Closure", level=2)
    if not vendors:
        _text_para(doc, "(none reported)", italic=True, color=MUTED)
        return
    t = doc.add_table(rows=1 + len(vendors), cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Vendor"
    hdr[1].text = "Contract Status"
    hdr[2].text = "Notes"
    for i, v in enumerate(vendors, 1):
        cells = t.rows[i].cells
        cells[0].text = str(v.get("vendor", "[TBD]"))
        cells[1].text = str(v.get("contract_status", "[TBD]"))
        cells[2].text = str(v.get("notes", ""))


def render_signoffs(doc, signoffs):
    _heading(doc, "Sign-off Block", level=2)
    if not signoffs:
        _text_para(doc, "[TBD — no sign-offs recorded]", italic=True)
        return
    t = doc.add_table(rows=1 + len(signoffs), cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, label in enumerate(["Name", "Role", "Decision", "Date"]):
        hdr[i].text = label
    for i, s in enumerate(signoffs, 1):
        cells = t.rows[i].cells
        cells[0].text = str(s.get("name", "[TBD]"))
        cells[1].text = str(s.get("role", ""))
        cells[2].text = str(s.get("decision", "[TBD]"))
        cells[3].text = str(s.get("date", ""))


def collect_gaps(data):
    gaps = []
    header = data.get("header", {}) or {}
    for k in ("project_name", "project_manager", "sponsor", "start_date", "closeout_date"):
        if is_tbd(header.get(k)):
            gaps.append(f"header.{k}: [TBD]")
    if is_tbd(data.get("executive_summary")):
        gaps.append("executive_summary: [TBD]")
    if not (data.get("scope_delivered") or []):
        gaps.append("scope_delivered: [TBD — empty]")
    cost = data.get("cost") or {}
    if cost.get("planned_usd") is None:
        gaps.append("cost.planned_usd: [TBD]")
    if cost.get("actual_usd") is None:
        gaps.append("cost.actual_usd: [TBD]")
    sched = data.get("schedule") or {}
    for k in ("planned_end", "actual_end"):
        if not sched.get(k):
            gaps.append(f"schedule.{k}: [TBD]")
    for d in data.get("deliverables") or []:
        if is_tbd(d.get("acceptance_status")) or d.get("acceptance_status") not in VALID_ACCEPTANCE:
            if d.get("acceptance_status") != "Pending":
                gaps.append(f"deliverable '{d.get('name', '?')}' acceptance: {d.get('acceptance_status', '[TBD]')}")
    if not (data.get("signoffs") or []):
        gaps.append("signoffs: [TBD — no sign-offs recorded]")
    return gaps


def render_appendix(doc, gaps):
    if not gaps:
        return
    _heading(doc, "Appendix — Draft Gaps", level=2)
    _text_para(doc, f"{len(gaps)} unresolved items — resolve before formal closeout.",
               italic=True, color=MUTED)
    for g in gaps:
        _bullet(doc, g)


def render(data, output_path: Path, questions_threshold=6):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    render_header(doc, data.get("header", {}) or {})
    render_summary(doc, data.get("executive_summary"))
    render_scope(doc, data.get("scope_delivered") or [], data.get("scope_deferred") or [])
    render_schedule(doc, data.get("schedule"))
    render_cost(doc, data.get("cost"))
    render_deliverables(doc, data.get("deliverables") or [])
    render_open_items(doc, data.get("open_items") or [])
    render_defects(doc, data.get("known_defects") or [])
    render_vendor_closure(doc, data.get("vendor_closure") or [])
    render_signoffs(doc, data.get("signoffs") or [])
    gaps = collect_gaps(data)
    render_appendix(doc, gaps)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if len(gaps) >= questions_threshold:
        oq = output_path.parent / "open_questions.md"
        lines = ["# Open Questions — Closeout Report", "",
                 f"This draft has **{len(gaps)} unresolved items**.", ""]
        for g in gaps:
            lines.append(f"- {g}")
        oq.write_text("\n".join(lines))
        return gaps, oq
    return gaps, None


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--input", help="Path to JSON input (omit to read stdin)")
    p.add_argument("--output", required=True)
    p.add_argument("--questions-threshold", type=int, default=6)
    args = p.parse_args()

    try:
        if args.input:
            data = json.loads(Path(args.input).read_text())
        else:
            data = json.loads(sys.stdin.read())
    except Exception as e:
        print(f"ERROR: could not parse input JSON: {e}", file=sys.stderr)
        return 2

    out = Path(args.output)
    gaps, oq = render(data, out, args.questions_threshold)
    print(f"Closeout report written: {out}")
    print(f"TBD count: {len(gaps)}")
    if oq:
        print(f"Open questions: {oq}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
